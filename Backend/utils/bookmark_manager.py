"""
Bookmark Manager - Handles detection and management of bookmarks in Word documents.
Provides utilities for auto-detecting bookmarks and suggesting/validating mappings.
"""

import logging
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import io

logger = logging.getLogger(__name__)


@dataclass
class BookmarkAnchor:
    """Reference information used to recreate a bookmark in another resume."""

    name: str
    paragraph_index: int
    paragraph_text: str
    heading_text: str
    relative_offset: int


class BookmarkManager:
    """Manages bookmarks in Word documents for point injection."""
    
    DEFAULT_SECTION_PATTERNS = {
        "PROFESSIONAL_SUMMARY": [
            "professional summary", "summary", "profile", "career summary", "objective"
        ],
        "SUMMARY": ["professional summary", "summary", "profile", "career summary", "objective"],
        "CORE_COMPETENCIES": [
            "core competencies", "technical skills", "skills", "technologies", "key skills"
        ],
        "SKILLS": ["core competencies", "technical skills", "skills", "technologies", "key skills"],
        "PROFESSIONAL_EXPERIENCE": [
            "professional experience", "work experience", "experience", "employment history",
            "career experience", "professional history"
        ],
        "EXPERIENCE": [
            "professional experience", "work experience", "experience", "employment history",
            "career experience", "professional history"
        ],
        "EDUCATION": ["education", "academic background", "academics"],
        "PROJECTS": ["projects", "project experience", "key projects", "professional projects"],
        "CERTIFICATIONS": ["certifications", "certificates", "licenses", "training"],
        "AWARDS": ["awards", "honors", "achievements", "recognition"],
        "LANGUAGES": ["languages", "language skills"],
    }
    
    def detect_bookmarks(self, resume_bytes) -> List[str]:
        """
        Detect all bookmarks in a Word document.
        
        Args:
            resume_bytes: BytesIO object of the resume document
            
        Returns:
            List of bookmark names found in the document
        """
        try:
            resume_bytes.seek(0)
            doc = Document(resume_bytes)
            resume_bytes.seek(0)
            
            bookmarks = []
            
            # Check document element for bookmarks
            for element in doc.element.iter():
                # Look for bookmark start elements
                if 'bookmarkStart' in element.tag:
                    bookmark_name = element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                    if bookmark_name:
                        bookmarks.append(bookmark_name)
            
            logger.debug(f"Detected bookmarks: {bookmarks}")
            return bookmarks
            
        except Exception as e:
            logger.error(f"Error detecting bookmarks: {e}")
            return []
    
    def ensure_bookmarks_from_reference(self, resume_bytes) -> Tuple[io.BytesIO, List[str], Dict]:
        """
        Ensure a resume has Word bookmarks. If it has none, recreate bookmarks by
        matching section locations from a reference resume that already has them.
        """
        resume_bytes.seek(0)
        existing = self.detect_bookmarks(resume_bytes)
        resume_bytes.seek(0)
        
        if existing:
            return resume_bytes, existing, {
                "auto_created": False,
                "created_count": 0,
                "reference_path": None,
                "matches": []
            }
        
        reference_path = self.find_reference_resume_path()
        if not reference_path:
            logger.warning("No bookmarked reference resume found for auto bookmark creation")
            return resume_bytes, [], {
                "auto_created": False,
                "created_count": 0,
                "reference_path": None,
                "matches": [],
                "message": "No bookmarked reference resume found"
            }
        
        try:
            with open(reference_path, "rb") as f:
                reference_bytes = io.BytesIO(f.read())
            anchors = self.extract_bookmark_anchors(reference_bytes)
        except Exception as e:
            logger.error(f"Could not read bookmark reference resume: {e}")
            return resume_bytes, [], {
                "auto_created": False,
                "created_count": 0,
                "reference_path": str(reference_path),
                "matches": [],
                "message": f"Could not read reference resume: {e}"
            }
        
        if not anchors:
            return resume_bytes, [], {
                "auto_created": False,
                "created_count": 0,
                "reference_path": str(reference_path),
                "matches": [],
                "message": "Reference resume has no usable bookmarks"
            }
        
        resume_bytes.seek(0)
        doc = Document(resume_bytes)
        created = []
        matches = []
        
        used_paragraphs = set()
        next_id = self._next_bookmark_id(doc)
        target_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        
        for anchor in anchors:
            matched_para, score, reason = self.find_matching_paragraph(anchor, target_paragraphs, used_paragraphs)
            if not matched_para and target_paragraphs:
                fallback_idx = min(anchor.paragraph_index, len(target_paragraphs) - 1)
                fallback_para = target_paragraphs[fallback_idx]
                if id(fallback_para) not in used_paragraphs:
                    matched_para = fallback_para
                    score = 45
                    reason = "reference-index"
            
            if not matched_para:
                logger.debug(f"No matching paragraph found for bookmark {anchor.name}")
                matches.append({"bookmark": anchor.name, "created": False, "score": 0, "reason": "no-match"})
                continue
            
            self.add_bookmark_to_paragraph(matched_para, anchor.name, next_id)
            next_id += 1
            used_paragraphs.add(id(matched_para))
            created.append(anchor.name)
            matches.append({
                "bookmark": anchor.name,
                "created": True,
                "score": score,
                "reason": reason,
                "target_text": matched_para.text.strip()[:120],
            })
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output, created, {
            "auto_created": bool(created),
            "created_count": len(created),
            "reference_path": str(reference_path),
            "matches": matches,
        }
    
    def find_reference_resume_path(self) -> Optional[Path]:
        """
        Find a DOCX resume with existing bookmarks to use as the location template.
        Prefer explicit env configuration, then scan the local resumes folder.
        """
        configured = (
            os.getenv("BOOKMARK_REFERENCE_RESUME_PATH")
            or os.getenv("BOOKMARK_TEMPLATE_PATH")
            or os.getenv("RESUME_BOOKMARK_TEMPLATE_PATH")
        )
        candidates = []
        
        if configured:
            candidates.append(Path(configured))
        
        candidates.extend(Path("./resumes").glob("*.docx"))
        candidates.extend(Path(".").glob("**/*template*.docx"))
        
        seen = set()
        for candidate in candidates:
            try:
                resolved = candidate.resolve()
            except Exception:
                resolved = candidate
            if resolved in seen or not candidate.exists() or candidate.suffix.lower() != ".docx":
                continue
            seen.add(resolved)
            
            try:
                with open(candidate, "rb") as f:
                    if self.detect_bookmarks(io.BytesIO(f.read())):
                        return candidate
            except Exception as e:
                logger.debug(f"Skipping bookmark reference candidate {candidate}: {e}")
        
        return None
    
    def extract_bookmark_anchors(self, resume_bytes) -> List[BookmarkAnchor]:
        """Extract bookmark names and nearby section context from a reference DOCX."""
        resume_bytes.seek(0)
        doc = Document(resume_bytes)
        anchors = []
        
        for idx, para in enumerate(doc.paragraphs):
            for bookmark_name in self._bookmark_names_in_paragraph(para):
                heading_idx, heading_text = self._nearest_heading_before(doc.paragraphs, idx)
                anchors.append(BookmarkAnchor(
                    name=bookmark_name,
                    paragraph_index=idx,
                    paragraph_text=para.text.strip(),
                    heading_text=heading_text,
                    relative_offset=idx - heading_idx if heading_idx is not None else 0,
                ))
        
        if anchors:
            return anchors
        
        # Fallback for bookmarks nested in XML that python-docx did not expose
        bookmarks = self.detect_bookmarks(resume_bytes)
        for bookmark_name in bookmarks:
            para = self.find_bookmark_paragraph(doc, bookmark_name)
            if para:
                idx = self._paragraph_index(doc.paragraphs, para)
                heading_idx, heading_text = self._nearest_heading_before(doc.paragraphs, idx)
                anchors.append(BookmarkAnchor(
                    name=bookmark_name,
                    paragraph_index=idx,
                    paragraph_text=para.text.strip(),
                    heading_text=heading_text,
                    relative_offset=idx - heading_idx if heading_idx is not None else 0,
                ))
        
        return anchors
    
    def find_bookmark_paragraph(self, doc: DocumentObject, bookmark_name: str) -> Optional[Paragraph]:
        """Find the paragraph that contains a bookmark."""
        for para in doc.paragraphs:
            if bookmark_name in para._element.xml:
                return para
        return None
    
    def find_matching_paragraph(
        self,
        anchor: BookmarkAnchor,
        target_paragraphs: List[Paragraph],
        used_paragraphs: set,
    ) -> Tuple[Optional[Paragraph], int, str]:
        """Find the best section-equivalent paragraph in a target resume."""
        best_para = None
        best_score = 0
        best_reason = "no-match"
        
        anchor_terms = self._bookmark_terms(anchor.name)
        source_text = self._normalize(anchor.paragraph_text)
        heading_text = self._normalize(anchor.heading_text)
        section_patterns = self._section_patterns_for(anchor.name)
        
        for idx, para in enumerate(target_paragraphs):
            if id(para) in used_paragraphs:
                continue
            
            text = para.text.strip()
            normalized = self._normalize(text)
            if not normalized:
                continue
            
            score = 0
            reasons = []
            
            if any(pattern and pattern in normalized for pattern in section_patterns):
                score += 85
                reasons.append("section-name")
            
            if heading_text and heading_text in normalized:
                score += 70
                reasons.append("reference-heading")
            
            if source_text and source_text in normalized:
                score += 65
                reasons.append("reference-text")
            
            matching_terms = [term for term in anchor_terms if term in normalized]
            if matching_terms:
                score += min(45, len(matching_terms) * 15)
                reasons.append("bookmark-terms")
            
            if self._looks_like_heading(para):
                score += 10
            
            # If the reference bookmark was a few paragraphs below a heading,
            # allow insertion near the equivalent heading instead of exact text.
            if score >= 75 and anchor.relative_offset > 0:
                offset_idx = min(idx + anchor.relative_offset, len(target_paragraphs) - 1)
                candidate = target_paragraphs[offset_idx]
                if id(candidate) not in used_paragraphs:
                    para = candidate
                    reasons.append("relative-offset")
            
            if score > best_score:
                best_para = para
                best_score = score
                best_reason = "+".join(reasons) if reasons else "low-confidence"
        
        if best_score < 45:
            return None, best_score, best_reason
        
        return best_para, best_score, best_reason
    
    def add_bookmark_to_paragraph(self, paragraph: Paragraph, bookmark_name: str, bookmark_id: int) -> None:
        """Insert Word bookmarkStart/bookmarkEnd XML around a paragraph's content."""
        if bookmark_name in paragraph._element.xml:
            return
        
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(bookmark_id))
        bookmark_start.set(qn("w:name"), bookmark_name)
        
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(bookmark_id))
        
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)
    
    def _bookmark_names_in_paragraph(self, paragraph: Paragraph) -> List[str]:
        names = []
        for element in paragraph._element.iter():
            if "bookmarkStart" in element.tag:
                name = element.get(qn("w:name"))
                if name:
                    names.append(name)
        return names
    
    def _next_bookmark_id(self, doc: DocumentObject) -> int:
        max_id = -1
        for element in doc.element.iter():
            if "bookmarkStart" in element.tag or "bookmarkEnd" in element.tag:
                value = element.get(qn("w:id"))
                if value and value.isdigit():
                    max_id = max(max_id, int(value))
        return max_id + 1
    
    def _paragraph_index(self, paragraphs: List[Paragraph], target: Paragraph) -> int:
        for idx, para in enumerate(paragraphs):
            if para._element == target._element:
                return idx
        return 0
    
    def _nearest_heading_before(self, paragraphs: List[Paragraph], index: int) -> Tuple[Optional[int], str]:
        for idx in range(index, -1, -1):
            para = paragraphs[idx]
            if self._looks_like_heading(para):
                return idx, para.text.strip()
        return None, ""
    
    def _looks_like_heading(self, paragraph: Paragraph) -> bool:
        text = paragraph.text.strip()
        if not text:
            return False
        if text.startswith(("•", "-", "*", "+")) or re.match(r"^\d+\.", text):
            return False
        
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name or "title" in style_name:
            return True
        
        normalized = self._normalize(text)
        known_sections = {
            pattern
            for patterns in self.DEFAULT_SECTION_PATTERNS.values()
            for pattern in patterns
        }
        if normalized in known_sections:
            return True
        
        word_count = len(text.split())
        has_sentence_punctuation = bool(re.search(r"[.!?]$", text))
        if text.isupper() and word_count <= 8 and not has_sentence_punctuation:
            return True
        
        return False
    
    def _section_patterns_for(self, bookmark_name: str) -> List[str]:
        normalized_name = re.sub(r"[^A-Za-z0-9]+", "_", bookmark_name).strip("_").upper()
        patterns = []
        for key, values in self.DEFAULT_SECTION_PATTERNS.items():
            if key in normalized_name or normalized_name in key:
                patterns.extend(values)
        
        patterns.extend(self._bookmark_terms(bookmark_name))
        return list(dict.fromkeys(self._normalize(pattern) for pattern in patterns if pattern))
    
    def _bookmark_terms(self, bookmark_name: str) -> List[str]:
        cleaned = re.sub(r"[^A-Za-z0-9]+", " ", bookmark_name)
        terms = [term for term in cleaned.lower().split() if len(term) > 2]
        if terms:
            terms.append(" ".join(terms))
        return terms
    
    def _normalize(self, value: str) -> str:
        return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9]+", " ", value.lower())).strip()
    
    def validate_mapping(self, cycle_to_bookmark: Dict[int, str], 
                        available_bookmarks: List[str]) -> Tuple[bool, str]:
        """
        Validate that a cycle-to-bookmark mapping is valid.
        
        Args:
            cycle_to_bookmark: Dict mapping cycle numbers to bookmark names
            available_bookmarks: List of available bookmarks in the document
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not cycle_to_bookmark:
            return False, "No mapping provided"
        
        if not available_bookmarks:
            return False, "No bookmarks available in document"
        
        # Check that all bookmark names in mapping exist
        for cycle_num, bookmark_name in cycle_to_bookmark.items():
            if bookmark_name not in available_bookmarks:
                return False, f"Bookmark '{bookmark_name}' not found in document"
        
        return True, "Mapping is valid"
    
    def suggest_mappings(self, available_bookmarks: List[str], 
                        num_cycles: int) -> Dict[int, str]:
        """
        Suggest an automatic mapping from cycles to bookmarks.
        
        Args:
            available_bookmarks: List of available bookmarks in the document
            num_cycles: Number of cycles to map
            
        Returns:
            Dict mapping cycle numbers to bookmark names
        """
        if not available_bookmarks:
            logger.warning("No bookmarks available for mapping")
            return {}
        
        mapping = {}
        
        # Simple approach: map cycles to bookmarks in order
        for cycle_num in range(1, num_cycles + 1):
            # Map cycle 1 to first bookmark, cycle 2 to second, etc.
            bookmark_idx = min(cycle_num - 1, len(available_bookmarks) - 1)
            mapping[cycle_num] = available_bookmarks[bookmark_idx]
        
        logger.debug(f"Suggested mapping: {mapping}")
        return mapping
